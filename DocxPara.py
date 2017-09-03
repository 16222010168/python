#coding:utf-8
from docx import Document
from xlwt import Workbook

document = Document('openSSH.docx')
paragraph = document.paragraphs
book = Workbook()
sheet1 = book.add_sheet('Sheet1')

ParaCount = len(paragraph)
#length = len(document.paragraphs)
l = []
for i in range(0,ParaCount):
    if document.paragraphs[i].text == '':
        pass
    elif document.paragraphs[i].text in u'影响IP清单：':
        pass
    else:
        l.append(document.paragraphs[i].text)

for i in range(0,len(l)):
    sheet1.write(0,i,l[i])

book.save('ssl.xls')

